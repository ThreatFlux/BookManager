"""
Compiler Module
--------------

Handles manuscript compilation using python-docx and WeasyPrint with configurable styling,
robust error handling, and progress tracking.

Features:
- Multiple output formats (PDF, DOCX)
- Configurable styling via config.yaml and environment variables
- Progress tracking with tqdm
- Comprehensive error handling and logging
- Support for various paper formats and style customization
"""

import os
import time
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, field

import markdown
from bs4 import BeautifulSoup
from docx import Document
from weasyprint import HTML, CSS
from weasyprint.text.fonts import FontConfiguration
from tqdm import tqdm

from ..utils.logging_setup import get_logger
from ..utils.config_loader import get_config

logger = get_logger(__name__)
__all__ = ["DocumentCompiler", "CompilationError", "compile_manuscript", "batch_compile"]


class CompilationError(Exception):
    """Custom exception for compilation errors."""


@dataclass
class PaperFormat:
    """
    Paper format configuration.

    Attributes:
        name: Standard paper format name (e.g., 'letter', 'a4')
        width: Page width with units
        height: Page height with units
    """

    name: str
    width: str
    height: str

    @classmethod
    def from_name(cls, name: str) -> "PaperFormat":
        """
        Create format from standard name.

        Args:
            name: Standard paper format name

        Returns:
            PaperFormat: Configured paper format
        """
        formats = {
            "letter": cls("letter", "8.5in", "11in"),
            "legal": cls("legal", "8.5in", "14in"),
            "a4": cls("a4", "210mm", "297mm"),
            "a5": cls("a5", "148mm", "210mm"),
        }
        return formats.get(name.lower(), formats["letter"])


@dataclass
class FontSettings:
    """
    Font-related settings.

    Attributes:
        body_font: Font family for body text
        heading_font: Font family for headings
        code_font: Font family for code blocks
        font_size: Base font size with units
    """

    body_font: str
    heading_font: str
    code_font: str
    font_size: str


@dataclass
class ColorSettings:
    """
    Color-related settings.

    Attributes:
        heading_color: Color for headings
        text_color: Color for body text
        link_color: Color for hyperlinks
        code_background: Background color for code blocks
    """

    heading_color: str
    text_color: str
    link_color: str
    code_background: str


@dataclass
class DocumentStyle:
    """
    Document styling configuration.

    Attributes:
        fonts: Font-related settings
        colors: Color-related settings
        paper_format: Paper size configuration
        margin_top: Top margin with units
        margin_right: Right margin with units
        margin_bottom: Bottom margin with units
        margin_left: Left margin with units
    """

    fonts: FontSettings
    colors: ColorSettings
    paper_format: PaperFormat = field(default_factory=lambda: PaperFormat.from_name("letter"))
    margin_top: str = field(default="1in")
    margin_right: str = field(default="1in")
    margin_bottom: str = field(default="1in")
    margin_left: str = field(default="1in")

    @classmethod
    def from_config(cls, config: Dict) -> "DocumentStyle":
        """
        Create style from config dictionary.

        Args:
            config: Configuration dictionary

        Returns:
            DocumentStyle: Configured document style
        """
        style_config = config.get("document_style", {})
        env_prefix = "BOOK_MANAGER_"

        fonts = FontSettings(
            body_font=os.getenv(
                f"{env_prefix}BODY_FONT",
                style_config.get("body_font", "'-apple-system', BlinkMacSystemFont, 'Segoe UI', Arial, sans-serif"),
            ),
            heading_font=os.getenv(
                f"{env_prefix}HEADING_FONT",
                style_config.get("heading_font", "'-apple-system', BlinkMacSystemFont, 'Segoe UI', Arial, sans-serif"),
            ),
            code_font=os.getenv(f"{env_prefix}CODE_FONT", style_config.get("code_font", "'Courier New', monospace")),
            font_size=os.getenv(f"{env_prefix}FONT_SIZE", style_config.get("font_size", "12pt")),
        )

        colors = ColorSettings(
            heading_color=os.getenv(f"{env_prefix}HEADING_COLOR", style_config.get("heading_color", "#000000")),
            text_color=os.getenv(f"{env_prefix}TEXT_COLOR", style_config.get("text_color", "#000000")),
            link_color=os.getenv(f"{env_prefix}LINK_COLOR", style_config.get("link_color", "#0366d6")),
            code_background=os.getenv(f"{env_prefix}CODE_BACKGROUND", style_config.get("code_background", "#f6f8fa")),
        )

        return cls(
            fonts=fonts,
            colors=colors,
            paper_format=PaperFormat.from_name(
                os.getenv(f"{env_prefix}PAPER_FORMAT", style_config.get("paper_format", "letter"))
            ),
            margin_top=os.getenv(f"{env_prefix}MARGIN_TOP", style_config.get("margin_top", "1in")),
            margin_right=os.getenv(f"{env_prefix}MARGIN_RIGHT", style_config.get("margin_right", "1in")),
            margin_bottom=os.getenv(f"{env_prefix}MARGIN_BOTTOM", style_config.get("margin_bottom", "1in")),
            margin_left=os.getenv(f"{env_prefix}MARGIN_LEFT", style_config.get("margin_left", "1in")),
        )


class ElementProcessor:
    """
    Handles processing of different HTML elements for DOCX conversion.

    Attributes:
        style: Document styling configuration
    """

    def __init__(self, style: DocumentStyle):
        """
        Initialize processor with document style.

        Args:
            style: Document styling configuration
        """
        self.style = style

    def process_heading(self, doc: Document, element: BeautifulSoup) -> None:
        """
        Process heading elements.

        Args:
            doc: Document being constructed
            element: BeautifulSoup heading element
        """
        level = int(element.name[1])
        paragraph = doc.add_paragraph(style=f"Heading {level}")
        text = element.get_text().strip()
        paragraph.add_run(text)

    def process_paragraph(self, doc: Document, element: BeautifulSoup) -> None:
        """
        Process paragraph elements.

        Args:
            doc: Document being constructed
            element: BeautifulSoup paragraph element
        """
        paragraph = doc.add_paragraph()
        self._process_text_with_formatting(paragraph, element)

    def process_list(self, doc: Document, element: BeautifulSoup) -> None:
        """
        Process list elements.

        Args:
            doc: Document being constructed
            element: BeautifulSoup list element
        """
        list_style = "List Bullet" if element.name == "ul" else "List Number"
        for li in element.find_all("li", recursive=False):
            paragraph = doc.add_paragraph(style=list_style)
            self._process_text_with_formatting(paragraph, li)

    def _process_text_with_formatting(self, paragraph, element) -> None:
        """
        Process text with formatting.

        Args:
            paragraph: Paragraph to add formatted text to
            element: BeautifulSoup element containing text
        """
        for child in element.children:
            if isinstance(child, str):
                if child.strip():
                    paragraph.add_run(child)
            elif child.name in ["strong", "b"]:
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name in ["em", "i"]:
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == "code":
                run = paragraph.add_run(child.get_text())
                run.font.name = self.style.fonts.code_font.split(",", maxsplit=1)[0].strip("'")


class DocumentCompiler:
    """
    Handles document compilation with configurable styling.

    Attributes:
        config: Configuration dictionary
        style: Document styling configuration
        supported_formats: List of supported output formats
        font_config: WeasyPrint font configuration
        element_processor: Processor for HTML elements
    """

    def __init__(self, config: Dict):
        """
        Initialize compiler with configuration.

        Args:
            config: Configuration dictionary
        """
        self.config = config
        self.style = DocumentStyle.from_config(config)
        self.supported_formats = ["pdf", "docx"]
        self.font_config = FontConfiguration()
        self.element_processor = ElementProcessor(self.style)

    def convert_to_docx(self, content: str, output_file: Path) -> None:
        """
        Convert markdown content to DOCX format.

        Args:
            content: Markdown content to convert
            output_file: Output file path

        Raises:
            CompilationError: If conversion fails
        """
        try:
            html_content = markdown.markdown(
                content, extensions=["fenced_code", "codehilite", "tables", "toc", "extra"]
            )

            doc = Document()
            soup = BeautifulSoup(html_content, "html.parser")

            # Process all elements
            for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "ul", "ol"]):
                if element.name.startswith("h"):
                    self.element_processor.process_heading(doc, element)
                elif element.name == "p":
                    self.element_processor.process_paragraph(doc, element)
                elif element.name in ["ul", "ol"]:
                    self.element_processor.process_list(doc, element)

            # Ensure output directory exists
            output_file.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_file))

        except (IOError, OSError) as e:
            raise CompilationError(f"File system error: {e}") from e
        except Exception as e:  # pylint: disable=broad-except
            raise CompilationError(f"DOCX conversion failed: {e}") from e

    def convert_to_pdf(self, content: str, output_file: Path) -> None:
        """
        Convert markdown content to PDF format.

        Args:
            content: Markdown content to convert
            output_file: Output file path

        Raises:
            CompilationError: If conversion fails
        """
        try:
            # Verify output directory exists
            if not output_file.parent.exists():
                raise CompilationError(f"Output directory does not exist: {output_file.parent}")

            html_content = markdown.markdown(content, extensions=["fenced_code", "codehilite", "tables", "toc"])

            # Create styled HTML
            styled_html = self._create_styled_html(html_content)
            html = HTML(string=styled_html)
            css = CSS(string=self._get_pdf_styles(), font_config=self.font_config)

            try:
                html.write_pdf(str(output_file), stylesheets=[css])
            except Exception as e:
                raise CompilationError(f"PDF writing failed: {e}") from e

        except (IOError, OSError) as e:
            raise CompilationError(f"File system error: {e}") from e
        except Exception as e:
            raise CompilationError(f"PDF conversion failed: {e}") from e

    def _create_styled_html(self, content: str) -> str:
        """
        Create HTML document with content.

        Args:
            content: HTML content

        Returns:
            str: Complete HTML document
        """
        return f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
        </head>
        <body>
            {content}
        </body>
        </html>
        """

    def _get_pdf_styles(self) -> str:
        """
        Get CSS styles based on configuration.

        Returns:
            str: CSS styles
        """
        return f"""
            @page {{
                margin: {self.style.margin_top} {self.style.margin_right} 
                        {self.style.margin_bottom} {self.style.margin_left};
                size: {self.style.paper_format.width} {self.style.paper_format.height};
                @top-right {{
                    content: counter(page);
                    font-family: {self.style.fonts.body_font};
                    font-size: {self.style.fonts.font_size};
                }}
            }}
            
            body {{
                font-family: {self.style.fonts.body_font};
                font-size: {self.style.fonts.font_size};
                line-height: 1.4;
                color: {self.style.colors.text_color};
                margin: 0;
                padding: 0;
            }}
            
            h1, h2, h3, h4, h5, h6 {{
                font-family: {self.style.fonts.heading_font};
                color: {self.style.colors.heading_color};
                margin-top: 1em;
                margin-bottom: 0.5em;
                border-bottom: 1px solid #eaecef;
                page-break-after: avoid;
            }}
            
            h1 {{ font-size: calc({self.style.fonts.font_size} * 2); }}
            h2 {{ font-size: calc({self.style.fonts.font_size} * 1.5); }}
            h3 {{ font-size: calc({self.style.fonts.font_size} * 1.3); }}
            
            p {{
                margin: 1em 0;
                orphans: 2;
                widows: 2;
            }}
            
            pre {{
                background-color: {self.style.colors.code_background};
                padding: 1em;
                margin: 1em 0;
                border-radius: 4px;
                white-space: pre-wrap;
                font-family: {self.style.fonts.code_font};
                font-size: calc({self.style.fonts.font_size} * 0.9);
            }}
            
            code {{
                background-color: {self.style.colors.code_background};
                padding: 0.2em 0.4em;
                border-radius: 3px;
                font-family: {self.style.fonts.code_font};
                font-size: calc({self.style.fonts.font_size} * 0.9);
            }}
            
            a {{
                color: {self.style.colors.link_color};
                text-decoration: none;
            }}
            
            ul, ol {{
                margin: 1em 0;
                padding-left: 2em;
            }}
            
            li {{
                margin: 0.5em 0;
            }}
            
            table {{
                border-collapse: collapse;
                width: 100%;
                margin: 1em 0;
            }}
            
            th, td {{
                border: 1px solid #dfe2e5;
                padding: 0.5em;
                text-align: left;
            }}
            
            thead {{
                background-color: {self.style.colors.code_background};
            }}
            
            img {{
                max-width: 100%;
                height: auto;
            }}
        """

    def compile_manuscript(self, content: str, formats: List[str], output_dir: Path) -> List[Path]:
        """
        Compile manuscript to specified formats.

        Args:
            content: Markdown content to compile
            formats: List of output formats
            output_dir: Output directory path

        Returns:
            List[Path]: List of generated files

        Raises:
            CompilationError: If compilation fails for any format
        """
        generated_files = []

        with tqdm(total=len(formats), desc="Compiling formats") as pbar:
            for fmt in formats:
                if fmt not in self.supported_formats:
                    logger.warning("Unsupported format: %s", fmt)
                    continue

                output_file = output_dir / f"manuscript.{fmt}"

                try:
                    if fmt == "docx":
                        self.convert_to_docx(content, output_file)
                    elif fmt == "pdf":
                        self.convert_to_pdf(content, output_file)

                    generated_files.append(output_file)

                except CompilationError as e:
                    logger.error("Failed to compile %s: %s", fmt, e)
                    raise

                pbar.update(1)

        return generated_files


def compile_manuscript(structure: Dict, formats: List[str], output_dir: Path, config: Dict) -> Tuple[bool, List[Path]]:
    """
    Compile manuscript from structure to specified formats.

    Args:
        structure: Book structure dictionary
        formats: List of output formats
        output_dir: Output directory path
        config: Configuration dictionary

    Returns:
        Tuple[bool, List[Path]]: Success status and list of generated files

    Raises:
        CompilationError: If compilation fails
    """
    compiler = DocumentCompiler(config)
    generated_files = []

    # Check for empty structure first
    if not structure:
        logger.warning("Empty book structure provided, nothing to compile")
        return False, []

    try:
        # Combine all scenes into single markdown
        content = []

        with tqdm(total=sum(len(acts) for acts in structure.values()), desc="Combining scenes") as pbar:
            for book_num in sorted(structure.keys()):
                content.append(f"\n# Book {book_num}\n")

                for act_num in sorted(structure[book_num].keys()):
                    content.append(f"\n## Act {act_num}\n")

                    for scene in sorted(structure[book_num][act_num], key=lambda x: x["scene_num"]):
                        scene_path = scene["path"]
                        content.append(f"\n### {scene_path.stem}\n")

                        try:
                            scene_content = Path(scene_path).read_text(encoding="utf-8")
                            content.append(scene_content)
                            content.append("\n")
                        except IOError as e:
                            logger.error("Error reading %s: %s", scene_path, e)
                            raise CompilationError(f"Failed to read scene: {e}") from e

                    pbar.update(1)

        # Compile combined content
        combined_content = "\n".join(content)

        # Don't compile if there's no actual content
        if not combined_content.strip():
            logger.warning("No content to compile")
            return False, []

        generated_files = compiler.compile_manuscript(combined_content, formats, output_dir)

        return len(generated_files) > 0, generated_files

    except CompilationError:
        raise
    except Exception as e:  # pylint: disable=broad-except
        logger.error("Unexpected error during compilation: %s", e)
        raise CompilationError(f"Unexpected compilation error: {e}") from e


def batch_compile(
    structure: Dict, formats: Optional[List[str]] = None, retries: int = 2, config: Optional[Dict] = None
) -> Tuple[bool, List[str]]:
    """
    Compile manuscript with retry logic.

    Args:
        structure: Book structure dictionary
        formats: List of output formats (default: from config)
        retries: Number of retry attempts
        config: Optional configuration override

    Returns:
        Tuple[bool, List[str]]: Success status and list of created files

    Example:
        >>> structure = {
        ...     1: {  # Book 1
        ...         1: [  # Act 1
        ...             {
        ...                 'path': 'scenes/scene1.md',
        ...                 'scene_num': 1
        ...             }
        ...         ]
        ...     }
        ... }
        >>> success, files = batch_compile(structure, formats=['pdf', 'docx'])
    """
    config = config or get_config()
    formats = formats or config.get("pandoc_output_formats", ["docx", "pdf"])
    output_dir = Path(config.get("compiled_dir", "Compiled"))
    created_files = []

    for attempt in range(retries + 1):
        if attempt > 0:
            logger.info("Retrying compilation %d/%d", attempt, retries)
            time.sleep(2**attempt)  # Exponential backoff

        try:
            success, files = compile_manuscript(structure, formats, output_dir, config)
            created_files.extend([str(f) for f in files])

            if success:
                return True, created_files

        except CompilationError as e:
            if attempt == retries:
                logger.error("Final compilation attempt failed: %s", e)
                raise
            logger.warning("Compilation failed, will retry: %s", e)

    return False, created_files


# Example usage
if __name__ == "__main__":
    import doctest

    doctest.testmod()
